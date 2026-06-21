# -*- coding: utf-8 -*-
"""Generate the Sampoorna SEO Product Requirements Document (.docx).

Run: python dev/make_prd.py
Output: docs/Sampoorna-SEO-PRD.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x1F, 0x4E, 0x79)      # deep blue
ACCENT = RGBColor(0x2E, 0x86, 0x3E)     # green
GREY = RGBColor(0x59, 0x59, 0x59)
LIGHT = "EAF1F8"

doc = Document()

# ---- Base styles -----------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)

for name, size, color in [("Heading 1", 16, BRAND), ("Heading 2", 13, BRAND),
                          ("Heading 3", 11.5, ACCENT)]:
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)


def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def para(text="", bold=False, italic=False, size=None, color=None, align=None,
         space_after=None, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        if size:
            r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    return p


def bullet(text, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def table(headers, rows, widths=None, header_fill=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    set_repeat_header(hdr)
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
        shade(c, header_fill or "1F4E79")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            pp = cells[i].paragraphs[0]
            run = pp.add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    para("", space_after=4)
    return t


def hrule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E79")
    pbdr.append(bottom)
    pPr.append(pbdr)


# ===========================================================================
# TITLE PAGE
# ===========================================================================
for _ in range(3):
    para("")
para("Product Requirements Document", bold=True, size=30, color=BRAND,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("Sampoorna SEO", bold=True, size=22, color=ACCENT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Agency-grade WordPress SEO engine + multi-site control plane + GEO/AI visibility",
     italic=True, size=12, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(6):
    para("")
para("Version 1.1  ·  21 June 2026", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Status: Draft for review", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Owner: Product / Engineering  ·  LSN Soft", size=11, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ===========================================================================
# DOCUMENT CONTROL
# ===========================================================================
doc.add_heading("Document Control", level=1)
table(["Version", "Date", "Author", "Summary of change"],
      [["0.1", "May 2026", "Product", "Initial requirements specification"],
       ["1.0", "20 Jun 2026", "Product / Eng", "PRD aligned to built Phases 0–5 (engine, schema, migration, control plane, GEO/AI)"],
       ["1.1", "21 Jun 2026", "Product / Eng", "Phase-4 breadth + depth shipped; added §14 Completion Status; scope/release refreshed"]],
      widths=[0.9, 1.2, 1.4, 3.3])

para("Related documents", bold=True, space_after=2)
bullet("Sampoorna SEO Requirements Specification v0.1 (source PRD)")
bullet("Sampoorna SEO — Build Plan & Program Roadmap v1.0")
bullet("WP Search Console Manager Plugin Spec (seed module)")
doc.add_page_break()

# ===========================================================================
# TABLE OF CONTENTS (static)
# ===========================================================================
doc.add_heading("Contents", level=1)
toc = [
    "1. Executive Summary", "2. Problem & Opportunity", "3. Goals & Success Metrics",
    "4. Target Users & Personas", "5. Product Scope", "6. System Architecture",
    "7. Functional Requirements", "8. Non-Functional Requirements",
    "9. Data, Security & Privacy", "10. Release Plan & Phasing",
    "11. Risks, Assumptions & Dependencies", "12. Open Questions", "13. Glossary",
    "14. Completion Status",
]
for item in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(item)
    r.font.size = Pt(11)
    if item[0].isdigit() and "." in item[:3]:
        r.bold = True
doc.add_page_break()

# ===========================================================================
# 1. EXECUTIVE SUMMARY
# ===========================================================================
doc.add_heading("1. Executive Summary", level=1)
para("Sampoorna SEO is an agency-grade SEO platform for WordPress. Unlike plugins that "
     "wrap an existing engine, Sampoorna builds the engine itself and renders all SEO "
     "output (meta, schema, canonicals, sitemaps) server-side into wp_head — more reliable "
     "for search engines and immune to the client-side-pixel failure modes of competitors "
     "such as OTTO.")
para("The product is two deployables working together:")
bullet("a per-site WordPress plugin (PHP 8.1+) that owns on-page SEO, schema, technical SEO, "
       "migration, and GEO/AI visibility; and", bold_lead="The engine — ")
bullet("a central control plane (Node.js + Fastify + PostgreSQL) that enrolls many sites over "
       "a signed handshake, scores their health, and runs an audit → approve → deploy → rollback "
       "workflow across the fleet from one dashboard.", bold_lead="The control plane — ")
para("Three capabilities define success and none is the on-page engine itself: the control "
     "plane (the only part that cannot be bought), GEO/AI visibility (llms.txt, AI-crawler "
     "access, AI-referral and LLM-citation tracking — a gap no WordPress-native plugin fills), "
     "and lossless migration (the gate for moving a live client without de-indexing them).")
para("As of this version, Phases 0–5 are built and verified: the engine, schema, lossless "
     "importers for Yoast / Rank Math / AIOSEO (posts and taxonomies), the full control-plane "
     "pipeline with authentication and scheduled metric collection, and the GEO/AI-visibility "
     "module set.", space_after=4)

# ===========================================================================
# 2. PROBLEM & OPPORTUNITY
# ===========================================================================
doc.add_heading("2. Problem & Opportunity", level=1)
doc.add_heading("2.1 Problem", level=2)
bullet("Agencies managing dozens of WordPress clients have no single place to configure, audit, "
       "approve, deploy, and roll back SEO changes across the whole fleet.")
bullet("Existing WP SEO plugins are single-site and offer no central control, no white-label "
       "operations, and no safe bulk deployment with rollback.")
bullet("Answer engines and LLMs (ChatGPT, Claude, Perplexity, Google AI) are becoming a major "
       "discovery surface, yet no WP-native tool measures or improves AI visibility "
       "(crawler access, citations, AI-referred traffic).")
bullet("Migrating a client off Yoast / Rank Math / AIOSEO risks losing meta, redirects, and "
       "schema — and therefore rankings — during cutover.")
doc.add_heading("2.2 Opportunity", level=2)
para("Because the vendor controls every client's WordPress install, Sampoorna can render SEO "
     "server-side and operate the fleet centrally — a structural advantage platform-agnostic "
     "competitors cannot match. The four competitive 'empty columns' are AI visibility, the "
     "control plane, white-label, and GEO/AEO; disproportionate effort goes there rather than "
     "to table-stakes on-page features.")

# ===========================================================================
# 3. GOALS & SUCCESS METRICS
# ===========================================================================
doc.add_heading("3. Goals & Success Metrics", level=1)
doc.add_heading("3.1 Product goals", level=2)
table(["#", "Goal"],
      [["G1", "Own a server-rendered SEO engine with parity on table-stakes on-page features."],
       ["G2", "Operate many client sites from one control plane with approve/deploy/rollback safety."],
       ["G3", "Lead on GEO/AI visibility: llms.txt, AI-crawler access, AI-referral & citation tracking."],
       ["G4", "Migrate clients losslessly from Yoast, Rank Math, and AIOSEO (posts + taxonomies)."],
       ["G5", "Keep AI keys and provider secrets off client sites; proxy all model calls centrally."]],
      widths=[0.5, 5.9])
doc.add_heading("3.2 Success metrics", level=2)
table(["Metric", "Target"],
      [["Migration fidelity", "100% of supported meta/redirects imported; zero de-index events at cutover"],
       ["Deploy safety", "Every deployed change reversible; rollback skips human-edited values"],
       ["Front-end performance", "SEO output served from cache; near-zero added TTFB; no wp_head fatals"],
       ["Fleet visibility", "5-dimension health score per site, refreshed on a schedule"],
       ["AI visibility", "Per-bot crawler activity, AI-referral traffic, and citation rate tracked per site"]],
      widths=[2.2, 4.2])

# ===========================================================================
# 4. PERSONAS
# ===========================================================================
doc.add_heading("4. Target Users & Personas", level=1)
table(["Persona", "Role", "Primary needs"],
      [["Agency operator", "Manages the control plane",
        "Enroll sites; see fleet health; approve and deploy changes; roll back safely; white-label."],
       ["Site SEO specialist", "Works in one site's WP admin",
        "Edit per-page meta/schema; run analysis; fix issues; migrate from a prior plugin."],
       ["Content editor", "Authors posts/pages",
        "Live on-page, readability, and AEO scoring while writing; clear, actionable guidance."],
       ["Client stakeholder", "Read-only viewer",
        "Branded reports and a trustworthy health/visibility summary."]],
      widths=[1.4, 1.7, 3.3])

# ===========================================================================
# 5. SCOPE
# ===========================================================================
doc.add_heading("5. Product Scope", level=1)
doc.add_heading("5.1 In scope (built)", level=2)
for t in [
    "Per-object SEO meta (title, description, focus keyword, robots, canonical) with templating and wp_head rendering",
    "Content scoring: on-page, readability, and Answer-Engine-Optimization (AEO), with live recompute in the editor",
    "JSON-LD schema graph: Article/WebPage/Organization/WebSite, BreadcrumbList, Person, plus FAQ, HowTo, Product (WooCommerce), and LocalBusiness",
    "Technical SEO: XML sitemaps, redirects, 404 monitoring, robots.txt control, IndexNow, Google Indexing API, Bing Webmaster submission",
    "Integrations: Google Search Console (OAuth + search-analytics) and GA4 traffic summary (Analytics Data API) on a shared Google connection",
    "Lossless migration importers: Yoast, Rank Math, AIOSEO — posts and term/taxonomy SEO — with dry-run",
    "Control plane: signed handshake, sites dashboard, 5-dimension health scoring (UX/CWV via PageSpeed), audit → approve → deploy → rollback",
    "Control-plane fleet ops: config templating by vertical, bulk operations, scheduled metrics/audits/citation, critical-event webhook alerts",
    "Control-plane access: authentication with admin / viewer / client roles (clients scoped to their own sites), white-label branding, branded printable reports",
    "GEO / AI visibility: llms.txt & llms-full.txt, AI-crawler access checker, AI-bot engagement logging, AI-referral tracking, LLM citation-tracking prototype",
]:
    bullet(t)
doc.add_heading("5.2 Out of scope (this version)", level=2)
for t in [
    "Production AI-content generation pipeline on live sites (the AI layer proxies through the control plane; unattended writes are gated)",
    "Google Merchant Center feed integration and Google Business Profile (GBP) automation / local-grid tracking",
    "Config-template precedence policy (template-vs-local override rules) and bulk deploy of audit findings across many sites at once",
    "Control-plane GEO/visibility rollups aggregated across the fleet",
    "Activation of credential-gated integrations is the operator's step (keys are not bundled): PageSpeed key, GA4 property, Indexing API service account, Bing API key",
]:
    bullet(t)

# ===========================================================================
# 6. ARCHITECTURE
# ===========================================================================
doc.add_heading("6. System Architecture", level=1)
para("Two deployables communicate over signed REST. The control plane is optional at runtime: "
     "a site degrades to local-only operation if the plane is unreachable, so a control-plane "
     "outage never breaks a client's front end.")
doc.add_heading("6.1 Per-site plugin (sampoorna-seo)", level=2)
para("Standard WordPress plugin, PHP 8.1+, PSR-4 autoloaded, organized as isolated opt-in "
     "modules behind a thin core. A broken module must never take down wp_head. Namespace "
     "Sampoorna\\SEO.", space_after=2)
table(["Module", "Responsibility"],
      [["Core", "Container, hooks, settings store, REST router, capabilities, DB versioning (dbDelta)"],
       ["Security", "AES-256-GCM crypto, signed-request signer/verifier, nonces"],
       ["Meta", "Per-object meta store, template engine, wp_head renderer"],
       ["Schema", "JSON-LD graph builder (Article, FAQ, Product, LocalBusiness)"],
       ["Technical", "Sitemaps, redirects, 404 monitor, robots, IndexNow"],
       ["Content", "On-page score, readability, AEO score"],
       ["Geo", "llms.txt, AI-crawler access checker, engagement & referral logging"],
       ["Integrations", "GSC (OAuth + search analytics), seeded from WPSCM"],
       ["Migration", "Source detectors + importers (Yoast / Rank Math / AIOSEO), dry-run/verify"],
       ["ControlPlane", "Handshake client, signed status/metrics/audit/apply/rollback, deploy journal"],
       ["Admin", "Editors, settings screens, list tables"]],
      widths=[1.4, 5.0])
doc.add_heading("6.2 Control plane (control-plane)", level=2)
para("Node.js + TypeScript + Fastify, PostgreSQL persistence, EJS server-side rendering. Uses "
     "only Node built-ins for crypto (HMAC / scrypt / AES-256-GCM) — no new runtime "
     "dependencies. Forward-only SQL migration runner. Signed-cookie sessions with admin/viewer roles.")
doc.add_heading("6.3 Signed handshake (contract v1)", level=2)
para("Every site↔plane call is HMAC-signed. Canonical string is METHOD \\n ROUTE \\n TIMESTAMP "
     "\\n sha256_hex(body); signature header sha256=<hmac>; headers X-Sampoorna-Key-Id / "
     "-Timestamp / -Signature; ±300s clock-skew window. Site secrets are AES-256-GCM encrypted "
     "at rest in the plane's vault and never leave it.")

# ===========================================================================
# 7. FUNCTIONAL REQUIREMENTS
# ===========================================================================
doc.add_heading("7. Functional Requirements", level=1)
para("Requirement IDs are grouped by module. Priority: M = must, S = should, C = could.",
     italic=True, color=GREY, space_after=4)

def freq(title, rows):
    doc.add_heading(title, level=2)
    table(["ID", "Requirement", "Pri"], rows, widths=[0.7, 5.2, 0.5])

freq("7.1 On-page meta & content", [
    ["MET-1", "Store and render title, meta description, focus keyword, robots, and canonical per post/page/term.", "M"],
    ["MET-2", "Support templating tokens (site name, separators, term/post fields) for titles & descriptions.", "M"],
    ["MET-3", "Render all meta server-side into wp_head without fighting redirect_canonical.", "M"],
    ["CON-1", "Compute on-page, readability, and AEO scores for the edited content.", "M"],
    ["CON-2", "Recompute scores live in the editor via AJAX as the author types.", "S"],
])
freq("7.2 Schema (JSON-LD)", [
    ["SCH-1", "Emit a connected @graph (Organization/WebSite/WebPage/Article) on relevant views.", "M"],
    ["SCH-2", "Add FAQ schema from structured FAQ content.", "S"],
    ["SCH-3", "Add Product schema with Offer and AggregateRating on WooCommerce product pages.", "S"],
    ["SCH-4", "Add LocalBusiness schema for multi-location businesses.", "S"],
])
freq("7.3 Technical SEO", [
    ["TEC-1", "Generate version-stamped, cached XML sitemaps served with near-zero TTFB.", "M"],
    ["TEC-2", "Manage redirects and monitor 404s (new/known states).", "M"],
    ["TEC-3", "Control robots.txt body and wp_robots output, including AI-crawler directives.", "M"],
    ["TEC-4", "Submit changes to IndexNow when enabled.", "S"],
])
freq("7.4 Integrations", [
    ["INT-1", "Connect Google Search Console via OAuth 2.0 (auth, callback, refresh, revoke).", "M"],
    ["INT-2", "Pull search-analytics insights and compare 28-day windows.", "S"],
    ["INT-3", "Relocate token custody to the control plane; per-site is a fallback.", "C"],
])
freq("7.5 Migration", [
    ["MIG-1", "Detect and import Yoast SEO meta for posts and taxonomies (incl. the wpseo_taxonomy_meta option).", "M"],
    ["MIG-2", "Import Rank Math meta (post meta + term meta) with token normalization.", "M"],
    ["MIG-3", "Import AIOSEO meta from its custom tables (posts + aioseo_terms).", "M"],
    ["MIG-4", "Provide a dry-run/verify mode before writing, and a WP-CLI command.", "S"],
])
freq("7.6 Control plane", [
    ["CP-1", "Enroll a site by key id + secret; verify all calls via the signed handshake.", "M"],
    ["CP-2", "Show a sites dashboard with online/stale/never status and last-seen.", "M"],
    ["CP-3", "Score each site on 5 dimensions (content, technical, authority, GEO, UX).", "M"],
    ["CP-4", "Run audit → approve → deploy → rollback; journal prior values; idempotent per deploy id.", "M"],
    ["CP-5", "Guard rollback so human-edited values are not overwritten; mark status only when restored.", "M"],
    ["CP-6", "Authenticate the dashboard (admin/viewer); viewers are read-only, non-GET requires admin.", "M"],
    ["CP-7", "Refresh all sites' metrics on a schedule (CP_REFRESH_MINUTES) and on demand.", "S"],
])
freq("7.7 GEO / AI visibility", [
    ["GEO-1", "Serve /llms.txt and /llms-full.txt (curated markdown map), cached and noindex.", "M"],
    ["GEO-2", "Check AI-crawler access (GPTBot, ClaudeBot, PerplexityBot, Google-Extended, CCBot) vs robots, with wildcard matching.", "M"],
    ["GEO-3", "Log AI-bot crawler engagement (per-bot frequency, last seen).", "S"],
    ["GEO-4", "Track AI-referral traffic from answer engines.", "S"],
    ["GEO-5", "Sample LLM citations (QUEST-style) with word-boundary brand matching; pluggable LLM client, deterministic stub by default.", "C"],
])

# ===========================================================================
# 8. NON-FUNCTIONAL
# ===========================================================================
doc.add_heading("8. Non-Functional Requirements", level=1)
table(["Area", "Requirement"],
      [["Performance", "All SEO output cached and version-stamped; near-zero added TTFB; no CDN dependency in production."],
       ["Reliability", "Modules isolated; a module fault must not break wp_head; control plane optional at runtime."],
       ["Security", "AES-256-GCM at rest; HMAC-signed requests with ±300s skew; AI/provider keys never on client sites."],
       ["Compatibility", "PHP 8.1+; WordPress current; WooCommerce-gated features no-op when WooCommerce is absent."],
       ["Maintainability", "PSR-4 + Composer (dev-only, no runtime third-party deps); phpcs/phpstan/phpunit green."],
       ["Observability", "Control-plane structured logs; scheduled-refresh run/total logged; site health snapshots persisted."]],
      widths=[1.5, 4.9])

# ===========================================================================
# 9. DATA, SECURITY & PRIVACY
# ===========================================================================
doc.add_heading("9. Data, Security & Privacy", level=1)
bullet("Trust boundary: AI keys and provider credentials live only in the control-plane vault; "
       "all model calls are proxied so cost and secrets stay with the vendor.", bold_lead="Secrets — ")
bullet("Site secrets are AES-256-GCM encrypted at rest; every site↔plane call is HMAC-signed and "
       "skew-bounded; the dashboard adds session auth with admin/viewer roles.", bold_lead="At rest & in transit — ")
bullet("Deploy journal records old and new values per change so any deployment is reversible; "
       "rollback skips values a human edited after deploy.", bold_lead="Reversibility — ")
bullet("Plugin persistence uses dbDelta with a tracked DB_VERSION; the control plane uses a "
       "forward-only SQL migration runner.", bold_lead="Schema management — ")

# ===========================================================================
# 10. RELEASE PLAN
# ===========================================================================
doc.add_heading("10. Release Plan & Phasing", level=1)
table(["Phase", "Scope", "Status"],
      [["0", "Foundation: rename to Sampoorna\\SEO, PSR-4 autoload, module loader, DB versioning", "Done"],
       ["1", "On-page meta store + wp_head renderer; content/readability/AEO scoring", "Done"],
       ["2", "Schema graph (Article, FAQ, Product, LocalBusiness) + live AJAX score recompute", "Done"],
       ["3", "Migration importers: Yoast, Rank Math, AIOSEO — posts + term/taxonomy", "Done"],
       ["4", "Control plane: handshake, sites dashboard, health scores, pipeline, auth, scheduled metrics", "Done"],
       ["4b", "Phase-4 breadth: config templating, bulk ops, white-label + client roles, alerts, branded reports", "Done"],
       ["5", "GEO/AI visibility: llms.txt, AI-crawler access, engagement, referral, citation prototype", "Done"],
       ["5b", "Depth: UX/CWV (PageSpeed), scheduled audits/citation, schema HowTo, GA4, Google Indexing API, Bing", "Done"],
       ["Next", "Operator activation of credential-gated integrations; Merchant Center / GBP; GEO rollups; bulk deploy", "Planned"]],
      widths=[0.8, 4.8, 0.8])

# ===========================================================================
# 11. RISKS
# ===========================================================================
doc.add_heading("11. Risks, Assumptions & Dependencies", level=1)
table(["Type", "Item", "Mitigation / note"],
      [["Risk", "Migration data loss at client cutover", "Dry-run/verify mode; lossless importers; per-source token maps"],
       ["Risk", "Engine scope creep ('tar pit')", "Strict per-module isolation; table-stakes parity, not gold-plating"],
       ["Risk", "Control-plane single point of failure", "Plane optional at runtime; sites degrade to local-only"],
       ["Risk", "Unattended AI output reaching live sites", "AI proxied through plane; writes gated; deterministic layer beneath"],
       ["Assumption", "Vendor controls each client's WordPress", "Enables server-side rendering advantage over OTTO"],
       ["Dependency", "Credential-gated integrations", "Built stub-first; operator supplies PageSpeed key, GA4 property, Indexing service account, Bing key, CP_LLM_KEY to activate"]],
      widths=[1.0, 2.6, 2.8])

# ===========================================================================
# 12. OPEN QUESTIONS
# ===========================================================================
doc.add_heading("12. Open Questions", level=1)
bullet("Config-template precedence: when a vertical template is pushed, may a site override locally, and does the next push re-assert? (Today: last-write-wins via the reversible deploy journal.)")
bullet("White-label depth for v2 — custom domains and full login/branding removal, beyond the current name/logo/accent/role scoping?")
bullet("Should bulk deploy of audit findings across many sites be added, and with what approval gate?")
bullet("Citation sampling cadence + per-client budget once a real CP_LLM_KEY is in use (scheduler is wired but defaults off).")

# ===========================================================================
# 13. GLOSSARY
# ===========================================================================
doc.add_heading("13. Glossary", level=1)
table(["Term", "Definition"],
      [["AEO", "Answer Engine Optimization — readiness of content to be quoted by answer engines."],
       ["GEO", "Generative Engine Optimization — visibility within LLM/AI answer surfaces."],
       ["Control plane", "Central Node service that enrolls, scores, and deploys to many sites."],
       ["Handshake", "HMAC-signed REST contract (v1) between a site and the control plane."],
       ["llms.txt", "Curated markdown map at /llms.txt telling LLMs what a site is and which pages matter."],
       ["Deploy journal", "Record of old/new values per change enabling guarded rollback."],
       ["TTFB", "Time To First Byte — front-end performance budget the SEO output must protect."]],
      widths=[1.3, 5.1])

# ===========================================================================
# 14. COMPLETION STATUS
# ===========================================================================
doc.add_heading("14. Completion Status", level=1)
para("As of v1.1 the product is feature-complete against the build plan. All six roadmap "
     "phases plus the Phase-4 breadth and post-launch depth items are built, verified "
     "(phpcs/phpstan/phpunit + control-plane typecheck/vitest, and live checks in Docker), "
     "and committed. The only remaining work to light up the credential-gated integrations is "
     "operator configuration — no code is pending for them.", space_after=4)

doc.add_heading("14.1 Shipped capabilities", level=2)
table(["Area", "Capability", "Status"],
      [["Engine", "Per-object meta + wp_head rendering, templating, on-page/readability/AEO scoring", "Done"],
       ["Schema", "@graph (Org/WebSite/WebPage/Article), BreadcrumbList, Person, FAQ, HowTo, Product, LocalBusiness", "Done"],
       ["Technical SEO", "Sitemaps, redirects, 404 monitor, robots, IndexNow, Google Indexing API, Bing submission", "Done"],
       ["Integrations", "Google Search Console (OAuth + analytics) and GA4 traffic summary", "Done"],
       ["Migration", "Yoast / Rank Math / AIOSEO importers — posts + taxonomies, dry-run", "Done"],
       ["Control plane — core", "Handshake, sites dashboard, 5-dim health (UX via PageSpeed), audit→approve→deploy→rollback", "Done"],
       ["Control plane — fleet ops", "Config templating, bulk operations, scheduled metrics/audits/citation, webhook alerts", "Done"],
       ["Control plane — access", "Admin / viewer / client roles, white-label branding, branded printable reports", "Done"],
       ["GEO / AI", "llms.txt, AI-crawler access checker, engagement + referral logging, citation prototype", "Done"]],
      widths=[1.5, 4.3, 0.6])

doc.add_heading("14.2 Needs operator activation (built, no code pending)", level=2)
bullet("UX/CWV health score — set a Google PageSpeed Insights API key (CP_PAGESPEED_KEY).")
bullet("GA4 traffic summary — connect Google (re-consent for the analytics scope) and set the GA4 property id.")
bullet("Google Indexing API — add a service-account JSON key (JobPosting / BroadcastEvent pages).")
bullet("Bing Webmaster submission — add a Bing Webmaster API key.")
bullet("Real LLM citation runs — set CP_LLM_KEY (defaults to the deterministic stub).")

doc.add_heading("14.3 Not yet built (optional / future)", level=2)
bullet("Merchant Center feed integration; Google Business Profile automation + local-grid tracking.")
bullet("Config-template precedence policy and bulk deploy of audit findings across many sites.")
bullet("Control-plane GEO/visibility rollups aggregated across the fleet.")

doc.add_heading("14.4 Verification baseline", level=2)
bullet("WordPress plugin: phpcs clean, phpstan clean, phpunit 146 tests / 507 assertions green.", bold_lead="PHP suite — ")
bullet("Control plane: tsc typecheck clean, vitest 39 tests green.", bold_lead="Node suite — ")
bullet("Each slice verified live in the Docker stack (WordPress + control plane) before commit.", bold_lead="Live — ")

hrule()
para("End of document — Sampoorna SEO Product Requirements Document v1.1",
     italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ---- Save ------------------------------------------------------------------
out_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "Sampoorna-SEO-PRD.docx")
doc.save(out_path)
print("WROTE", out_path)
